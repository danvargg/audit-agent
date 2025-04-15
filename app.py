import streamlit as st
import requests
import pandas as pd
from io import BytesIO
from docx import Document

st.title("ISO 13485:2016 Compliance Auditor")

st.sidebar.header("Upload Files")
standard_file = st.sidebar.file_uploader("Upload ISO Standard PDF", type=["pdf"])
document_file = st.sidebar.file_uploader("Upload Document PDF", type=["pdf"])

if st.sidebar.button("Analyze"):
    if standard_file and document_file:
        with st.spinner("Analyzing..."):
            files = {
                "standard": ("standard.pdf", standard_file, "application/pdf"),
                "document": ("document.pdf", document_file, "application/pdf"),
            }
            response = requests.post("http://127.0.0.1:8000/analyze", files=files)

            if response.status_code == 200:
                findings = response.json().get("findings", "No findings returned.")
                st.success("Analysis Complete!")

                # Parse findings into a DataFrame (assuming findings are structured)
                data = [
                    {
                        "Document Section": "7.1 Design and Development Planning",
                        "Finding": "No clear `trigger conditions` for updating the design and development plan.",
                        "ISO Standard Reference": "7.3.2",
                    },
                    {
                        "Document Section": "7.3 Design Input Requirements",
                        "Finding": "No explicit mention of a traceability matrix linking design inputs to verification steps.",
                        "ISO Standard Reference": "7.3.3",
                    },
                    # Add more rows based on findings
                ]
                df = pd.DataFrame(data)

                # Display summary table
                st.table(df)

                # Generate Word document
                def generate_word_doc(data):
                    doc = Document()
                    doc.add_heading("Audit Findings Report", level=1)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Document Section"
                    hdr_cells[1].text = "Finding"
                    hdr_cells[2].text = "ISO Standard Reference"
                    for row in data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = row["Document Section"]
                        row_cells[1].text = row["Finding"]
                        row_cells[2].text = row["ISO Standard Reference"]
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer

                # Add download button
                word_doc = generate_word_doc(data)
                st.download_button(
                    label="Download Findings as Word Document",
                    data=word_doc,
                    file_name="audit_findings.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error(f"Error: {response.json().get('error', 'Unknown error')}")
    else:
        st.warning("Please upload both files.")

